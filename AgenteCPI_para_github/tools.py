"""
Lógica de negocio compartida entre mcp_server.py y agent.py.
Ambos importan de acá para no duplicar código.
"""
import json
import os
import re
import base64
import zipfile
import io
import xml.etree.ElementTree as ET
import requests as _requests
from datetime import datetime
from dotenv import load_dotenv
from cpi_client import CPIClient

load_dotenv(override=True)

_client: CPIClient | None = None


def get_client() -> CPIClient:
    global _client
    if _client is None:
        _client = CPIClient()
    return _client


# ------------------------------------------------------------------
# Tool: listar paquetes
# ------------------------------------------------------------------
def tool_list_packages(filter: str = "") -> str:
    """
    Lista los Integration Packages del tenant.
    Si se pasa filter, filtra por nombre o ID (case-insensitive).
    """
    packages = get_client().filter_packages(filter)
    result = [
        {
            "id":          p.get("Id", ""),
            "name":        p.get("Name", ""),
            "description": p.get("Description", ""),
            "version":     p.get("Version", ""),
        }
        for p in packages
    ]
    return json.dumps(result, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------
# Tool: listar iFlows filtrando por paquete
# ------------------------------------------------------------------
def tool_list_iflows(package_filter: str = "") -> str:
    """
    Lista iFlows de los paquetes que coincidan con package_filter.
    Sin filtro devuelve todos los iFlows del tenant.
    """
    iflows = get_client().filter_iflows(package_filter)
    result = [
        {
            "id":       f.get("Id", ""),
            "name":     f.get("Name", ""),
            "version":  f.get("Version", ""),
            "package":  f.get("_PackageName", ""),
            "pkg_id":   f.get("_PackageId", ""),
        }
        for f in iflows
    ]
    return json.dumps(result, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------
# Tool: iFlows de un paquete específico
# ------------------------------------------------------------------
def tool_get_iflows_for_package(package_id: str) -> str:
    """
    Devuelve todos los iFlows de un paquete dado su ID exacto.
    """
    iflows = get_client().get_iflows_for_package(package_id)
    result = [
        {
            "id":      f.get("Id", ""),
            "name":    f.get("Name", ""),
            "version": f.get("Version", ""),
        }
        for f in iflows
    ]
    return json.dumps(result, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------
# Tool: backup de iFlow a GitHub
# ------------------------------------------------------------------
def tool_backup_iflow(iflow_id: str) -> str:
    """
    Descarga un iFlow de SAP CPI y lo sube a GitHub como backup.
    - Crea una carpeta con el ID del iFlow en el repo.
    - El archivo se nombra: {iflow_id}_{YYYYMMDD_HHMMSS}.zip
    """
    github_token = os.getenv("GITHUB_TOKEN")
    github_repo  = os.getenv("GITHUB_REPO", "cdrrodriguez/CPI-iflows-ejemplos")

    if not github_token:
        return json.dumps({"error": "GITHUB_TOKEN no configurado en .env"})

    # 1. Descargar el iFlow desde CPI
    try:
        zip_bytes = get_client().download_iflow(iflow_id)
    except Exception as e:
        return json.dumps({"error": f"Error al descargar iFlow '{iflow_id}': {str(e)}"})

    # 2. Armar nombre y path del archivo
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"{iflow_id}_{timestamp}.zip"
    path      = f"{iflow_id}/{filename}"

    # 3. Subir a GitHub via API
    content_b64 = base64.b64encode(zip_bytes).decode("utf-8")
    url = f"https://api.github.com/repos/{github_repo}/contents/{path}"
    headers = {
        "Authorization": f"token {github_token}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json",
    }
    body = {
        "message": f"Backup iFlow {iflow_id} - {timestamp}",
        "content": content_b64,
    }

    try:
        resp = _requests.put(url, headers=headers, json=body, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        html_url = data.get("content", {}).get("html_url", "")
        return json.dumps({
            "status":      "ok",
            "iflow_id":    iflow_id,
            "filename":    filename,
            "path":        path,
            "github_url":  html_url,
            "size_bytes":  len(zip_bytes),
            "timestamp":   timestamp,
            "repo":        github_repo,
        }, ensure_ascii=False)
    except _requests.HTTPError as e:
        detail = ""
        try:
            detail = e.response.json().get("message", e.response.text[:300])
        except Exception:
            pass
        return json.dumps({"error": f"Error al subir a GitHub: {str(e)}", "detail": detail})


# ------------------------------------------------------------------
# Tool: analizar contenido de un iFlow
# ------------------------------------------------------------------
_TEXT_EXTENSIONS = (".iflw", ".groovy", ".js", ".xsl", ".xslt",
                    ".xml", ".prop", ".mf", ".json", ".properties")
_MAX_FILE_BYTES  = 80_000   # ~80 KB por archivo para no saturar el contexto


def tool_analyze_iflow(iflow_id: str) -> str:
    """
    Descarga un iFlow de SAP CPI, extrae su contenido y devuelve los archivos
    relevantes (flujo BPMN, scripts, mappings, configuración) para que el agente
    pueda explicar qué hace, qué APIs usa y cómo se autentica.
    """
    # 1. Descargar el ZIP
    try:
        zip_bytes = get_client().download_iflow(iflow_id)
    except Exception as e:
        return json.dumps({"error": f"Error al descargar iFlow '{iflow_id}': {str(e)}"})

    # 2. Extraer archivos relevantes en memoria
    files_content  = {}   # { filename: text_content }
    files_skipped  = []   # archivos binarios / muy grandes
    all_files      = []

    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            all_files = zf.namelist()
            for name in all_files:
                lower = name.lower()
                # Ignorar directorios
                if name.endswith("/"):
                    continue
                # Solo archivos de texto relevantes
                if any(lower.endswith(ext) for ext in _TEXT_EXTENSIONS):
                    try:
                        raw = zf.read(name)
                        if len(raw) > _MAX_FILE_BYTES:
                            # Incluir solo el inicio si es muy grande
                            files_content[name] = (
                                raw[:_MAX_FILE_BYTES].decode("utf-8", errors="replace")
                                + f"\n... [TRUNCADO — archivo de {len(raw)} bytes]"
                            )
                        else:
                            files_content[name] = raw.decode("utf-8", errors="replace")
                    except Exception:
                        files_skipped.append(name)
                else:
                    files_skipped.append(name)
    except Exception as e:
        return json.dumps({"error": f"Error al leer ZIP: {str(e)}"})

    # 3. Clasificar archivos por tipo para facilitar el análisis
    flow_def   = {k: v for k, v in files_content.items() if k.endswith(".iflw")}
    scripts    = {k: v for k, v in files_content.items()
                  if k.endswith(".groovy") or k.endswith(".js")}
    mappings   = {k: v for k, v in files_content.items()
                  if k.endswith(".xsl") or k.endswith(".xslt")}
    config     = {k: v for k, v in files_content.items()
                  if k not in {**flow_def, **scripts, **mappings}}

    # Enriquecer con iFlows similares del tenant (RAG)
    similar_iflows = _get_rag_context(iflow_id, exclude_id=iflow_id, n=2)

    return json.dumps({
        "iflow_id":        iflow_id,
        "zip_size_bytes":  len(zip_bytes),
        "total_files":     len(all_files),
        "files_list":      all_files,
        "flow_definition": flow_def,    # XML BPMN principal — acá están adapters y auth
        "scripts":         scripts,     # Groovy / JS
        "mappings":        mappings,    # XSLT
        "config_files":    config,      # MANIFEST, .prop, .properties
        "files_skipped":   files_skipped,
        "rag_contexto":    similar_iflows,
    }, ensure_ascii=False)


# ------------------------------------------------------------------
# Helpers para generación de Word
# ------------------------------------------------------------------
def _add_cell_shading(cell, fill_hex: str):
    """Aplica color de fondo a una celda de tabla en python-docx."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _add_hr(doc, color: str = "CCCCCC"):
    """Agrega una línea horizontal al documento."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _parse_inline(paragraph, text: str):
    """Agrega runs con formato inline (**bold**, *italic*, `code`) a un párrafo docx."""
    from docx.shared import Pt, RGBColor
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


# ------------------------------------------------------------------
# Tool: generar documento Word con la documentación del iFlow
# ------------------------------------------------------------------
def tool_document_iflow(iflow_id: str, analysis_markdown: str = "") -> str:
    """
    Genera un documento Word (.docx) con la documentación técnica de un iFlow.
    Descarga el iFlow, genera el análisis con Claude y crea el DOCX.
    Guarda el archivo en downloads/ y retorna la URL de descarga.
    """
    # Si no viene el análisis, generarlo automáticamente
    if not analysis_markdown.strip():
        import anthropic
        api_key = os.getenv("ANTHROPIC_API_KEY")
        claude  = anthropic.Anthropic(api_key=api_key)

        raw_json = tool_analyze_iflow(iflow_id)
        prompt = (
            f"Analizá el siguiente iFlow de SAP CPI y generá documentación técnica completa en Markdown.\n"
            f"Usá estas secciones: ## Resumen, ## Propósito, ## Flujo de pasos, "
            f"## Sistemas conectados, ## Autenticación, ## Transformaciones y mappings, "
            f"## Scripts, ## Parámetros externalizables, ## Diagrama de flujo.\n"
            f"Sé detallado y técnico. Respondé solo con el Markdown, sin texto previo.\n\n"
            f"Contenido del iFlow:\n{raw_json[:60000]}"
        )
        resp = claude.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        analysis_markdown = resp.content[0].text
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return json.dumps({"error": "python-docx no instalado. Ejecutar: pip install python-docx"})

    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Estilos de heading con color SAP ──────────────────────────────
    SAP_DARK  = RGBColor(0x00, 0x3B, 0x62)
    SAP_MID   = RGBColor(0x00, 0x5A, 0x94)
    for name, size, color in [
        ("Heading 1", 16, SAP_DARK),
        ("Heading 2", 13, SAP_DARK),
        ("Heading 3", 11, SAP_MID),
    ]:
        try:
            s = doc.styles[name]
            s.font.color.rgb = color
            s.font.size      = Pt(size)
            s.font.bold      = True
        except Exception:
            pass

    # ── Portada ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SAP CPI  ·  Documentación de iFlow")
    run.font.size      = Pt(10)
    run.font.color.rgb = SAP_DARK
    run.font.bold      = True

    h = doc.add_heading(iflow_id, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}    "
        f"Repo: cdrrodriguez/CPI-iflows-ejemplos"
    )
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _add_hr(doc, "003B62")
    doc.add_paragraph()

    # ── Parser de Markdown → elementos docx ──────────────────────────
    lines        = analysis_markdown.split("\n")
    i            = 0
    in_code      = False
    code_lines   = []

    while i < len(lines):
        line    = lines[i]
        stripped = line.rstrip()

        # ── Bloque de código ──────────────────────────────────────────
        if stripped.lstrip().startswith("```"):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    p   = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run("\n".join(code_lines))
                    run.font.name      = "Courier New"
                    run.font.size      = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x1F, 0x50, 0x80)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        s = stripped.strip()

        # ── Tabla Markdown ────────────────────────────────────────────
        if s.startswith("|") and "|" in s[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = [h.strip().strip("*`") for h in table_lines[0].split("|")[1:-1]]
                data_rows = []
                for tl in table_lines[2:]:
                    if re.match(r"^[\|\-\:\s]+$", tl):
                        continue
                    row = [c.strip() for c in tl.split("|")[1:-1]]
                    if any(c.strip() for c in row):
                        data_rows.append(row)

                if headers and data_rows:
                    ncols = len(headers)
                    tbl   = doc.add_table(rows=1 + len(data_rows), cols=ncols)
                    tbl.style = "Table Grid"

                    # Header row con fondo SAP
                    for j, hdr in enumerate(headers):
                        if j >= ncols:
                            break
                        cell = tbl.rows[0].cells[j]
                        cell.text = hdr
                        run = cell.paragraphs[0].runs[0]
                        run.bold            = True
                        run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
                        _add_cell_shading(cell, "003B62")

                    # Filas de datos
                    for ri, row in enumerate(data_rows):
                        row_cells = tbl.rows[ri + 1].cells
                        fill = "EEF4FB" if ri % 2 == 1 else "FFFFFF"
                        for j, ct in enumerate(row):
                            if j >= ncols:
                                break
                            clean = re.sub(r"\*\*|`", "", ct)
                            row_cells[j].text = clean
                            _add_cell_shading(row_cells[j], fill)

                    doc.add_paragraph()
            continue

        # ── Headings ──────────────────────────────────────────────────
        if s.startswith("### "):
            doc.add_heading(s[4:].strip("#").strip(), level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip("#").strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip("#").strip(), level=1)

        # ── Separador horizontal ──────────────────────────────────────
        elif re.match(r"^[-_*]{3,}$", s):
            _add_hr(doc)

        # ── Lista bullet ──────────────────────────────────────────────
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline(p, s[2:])

        # ── Lista numerada ────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, re.sub(r"^\d+\.\s", "", s))

        # ── Cita / blockquote ─────────────────────────────────────────
        elif s.startswith("> "):
            p   = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(s[2:])
            run.italic          = True
            run.font.color.rgb  = RGBColor(0x55, 0x55, 0x55)

        # ── Línea vacía ───────────────────────────────────────────────
        elif not s:
            pass  # saltar; el espaciado viene del estilo del párrafo siguiente

        # ── Párrafo normal ────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            _parse_inline(p, s)

        i += 1

    # ── Guardar archivo ───────────────────────────────────────────────
    downloads_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "downloads")
    os.makedirs(downloads_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_id   = re.sub(r"[^\w\-.]", "_", iflow_id)[:80]
    filename  = f"{safe_id}_{timestamp}.docx"
    filepath  = os.path.join(downloads_dir, filename)
    doc.save(filepath)

    return json.dumps({
        "status":        "ok",
        "filename":      filename,
        "download_url":  f"/download/{filename}",
        "iflow_id":      iflow_id,
        "timestamp":     timestamp,
        "size_bytes":    os.path.getsize(filepath),
    }, ensure_ascii=False)


# ------------------------------------------------------------------
# Helpers de parseo XML para get_iflow_profile
# ------------------------------------------------------------------
def _xtag(el) -> str:
    """Retorna el tag local de un elemento XML ignorando el namespace."""
    t = el.tag
    return t.split("}")[-1] if "}" in t else t


def _read_props(extension_el) -> dict:
    """
    Lee las propiedades de un <extensionElements>.
    SAP CPI usa DOS formatos según la versión del iFlow:
      Formato A (atributos):  <ifl:property key="K" value="V"/>
      Formato B (elementos):  <ifl:property><key>K</key><value>V</value></ifl:property>
    """
    props = {}
    for prop in extension_el:
        if _xtag(prop) != "property":
            continue
        # Formato A — atributos
        k = prop.get("key", "")
        v = prop.get("value", "")
        if k:
            props[k] = v
            continue
        # Formato B — elementos hijo <key> / <value>
        k_el = next((c for c in prop if _xtag(c) == "key"),   None)
        v_el = next((c for c in prop if _xtag(c) == "value"), None)
        if k_el is not None:
            props[k_el.text or ""] = (v_el.text or "") if v_el is not None else ""
    return props


def _collect_steps(root) -> list[dict]:
    """
    Recorre el árbol XML y recolecta pasos con sus propiedades.
    Retorna lista de {element, name, props: {key: value}}.
    """
    STEP_TAGS = {"participant", "serviceTask", "callActivity",
                 "subProcess", "userTask", "scriptTask"}
    steps = []
    for el in root.iter():
        if _xtag(el) not in STEP_TAGS:
            continue
        name  = el.get("name", "")
        props = {}
        for child in el:
            if _xtag(child) == "extensionElements":
                props = _read_props(child)
                break
        if props:
            steps.append({"element": _xtag(el), "name": name, "props": props})
    return steps


# ------------------------------------------------------------------
# Tool: perfil técnico estructurado de un iFlow
# ------------------------------------------------------------------
def tool_get_iflow_profile(iflow_id: str) -> str:
    """
    Descarga un iFlow y parsea su XML para devolver un perfil técnico estructurado:
      - tipo_integracion : síncrona / asíncrona
      - endpoints        : URLs y rutas externas usadas
      - adapters         : tipos de adapter (SOAP, HTTP, OData, ProcessDirect, etc.)
      - mappings         : cantidad y complejidad estimada
      - scripts          : archivos Groovy/JS con métricas de complejidad
      - dependencias     : llamadas a otros iFlows (ProcessDirect) o servicios externos

    NOTA: En SAP CPI, el tipo de adapter está en messageFlow[name] (SOAP, HTTP, etc.),
    NO en la propiedad ComponentType del serviceTask. Los valores de address suelen ser
    parámetros externalizables (ej: {{WSFEAddress}}) o strings vacíos.
    """
    # 1. Descargar ZIP ------------------------------------------------
    try:
        zip_bytes = get_client().download_iflow(iflow_id)
    except Exception as e:
        return json.dumps({"error": f"No se pudo descargar el iFlow: {e}"})

    # 2. Extraer archivos relevantes ----------------------------------
    iflw_xml      = None
    script_files  = []
    mapping_files = []

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        for name in zf.namelist():
            lower = name.lower()
            if lower.endswith(".iflw"):
                iflw_xml = zf.read(name).decode("utf-8", errors="replace")
            elif lower.endswith(".groovy") or (lower.endswith(".js") and not lower.endswith(".json")):
                content = zf.read(name).decode("utf-8", errors="replace")
                script_files.append({"nombre": name.split("/")[-1], "contenido": content})
            elif lower.endswith((".mmap", ".xsl", ".xslt")):
                mapping_files.append(name.split("/")[-1])

    if not iflw_xml:
        return json.dumps({"error": "No se encontró el archivo .iflw en el ZIP"})

    # 3. Parsear XML --------------------------------------------------
    try:
        root = ET.fromstring(iflw_xml.encode("utf-8"))
    except ET.ParseError as e:
        return json.dumps({"error": f"Error al parsear XML: {e}"})

    # 4. Extraer adapters desde messageFlow ---------------------------
    # En SAP CPI el tipo de adapter está en messageFlow[name]
    # Las propiedades de config están en sus extensionElements
    _ADDR_KEYS = ("address", "private.address", "httpAddress",
                  "Private.httpAddress", "wsdlUrl", "serviceUrl", "targetUrl",
                  "wssdUrl", "soapWsdlURL")
    _ASYNC_ADAPTERS = {"JMS", "AMQP", "AdvancedEventMesh", "Kafka",
                       "AzureServiceBus", "SQS", "XI"}
    _INTERNAL_ADAPTERS = {"ProcessDirect", "script", "mapping", "contentEnricher",
                          "contentModifier", "router", "aggregator", "splitter",
                          "RequestReply", "SendStep"}

    adapters_raw   = []
    dependencies   = []
    async_signals  = []
    seen_dep_keys  = set()

    for mf in root.iter():
        if _xtag(mf) != "messageFlow":
            continue

        adapter_name = mf.get("name", "").strip()
        if not adapter_name:
            continue

        # Leer props del extensionElements del messageFlow
        props = {}
        for child in mf:
            if _xtag(child) == "extensionElements":
                props = _read_props(child)
                break

        # Intentar obtener la dirección (puede ser un param externalizable)
        addr = next((props[k] for k in _ADDR_KEYS if props.get(k, "").strip()), "")

        # Determinar si es sender o receiver según sourceRef / targetRef
        source_ref = mf.get("sourceRef", "")
        target_ref = mf.get("targetRef", "")

        adapters_raw.append({
            "tipo":      adapter_name,
            "address":   addr,
            "sourceRef": source_ref,
            "targetRef": target_ref,
            "props":     props,
        })

        # Dependencias ProcessDirect
        if adapter_name == "ProcessDirect" and addr:
            dep_key = f"pd:{addr}"
            if dep_key not in seen_dep_keys:
                seen_dep_keys.add(dep_key)
                dependencies.append({
                    "tipo":      "iFlow interno (ProcessDirect via messageFlow)",
                    "direccion": addr,
                })

        # Señales de asincronismo
        if adapter_name in _ASYNC_ADAPTERS:
            async_signals.append(f"adapter {adapter_name}")
        if props.get("isAsynchronous", "").lower() in ("true", "1", "yes"):
            async_signals.append("propiedad isAsynchronous=true")

    # 5. Dependencias adicionales desde callActivity ------------------
    # callActivity con subActivityType=NonLoopingProcess = llamada a otro iFlow
    for ca in root.iter():
        if _xtag(ca) != "callActivity":
            continue

        props = {}
        for child in ca:
            if _xtag(child) == "extensionElements":
                props = _read_props(child)
                break

        sub_type   = props.get("subActivityType", "").strip()
        process_id = props.get("processId", "").strip()
        ca_name    = ca.get("name", "").strip()

        if sub_type == "NonLoopingProcess" and process_id:
            dep_key = f"ca:{process_id}"
            if dep_key not in seen_dep_keys:
                seen_dep_keys.add(dep_key)
                dependencies.append({
                    "tipo":     "iFlow interno (callActivity)",
                    "iflow_id": process_id,
                    "step":     ca_name,
                })
        elif sub_type == "NonLoopingProcess" and ca_name:
            dep_key = f"ca_name:{ca_name}"
            if dep_key not in seen_dep_keys:
                seen_dep_keys.add(dep_key)
                dependencies.append({
                    "tipo": "iFlow interno (callActivity)",
                    "step": ca_name,
                })

        # ComponentType ProcessDirect en callActivity props
        comp = props.get("ComponentType", "").strip()
        if comp == "ProcessDirect":
            addr = props.get("address", "").strip()
            dep_key = f"pd_ca:{addr or ca_name}"
            if dep_key not in seen_dep_keys:
                seen_dep_keys.add(dep_key)
                dependencies.append({
                    "tipo":      "iFlow interno (ProcessDirect via callActivity)",
                    "direccion": addr,
                    "step":      ca_name,
                })

    # 6. Detectar mappings en el flujo (serviceTask activityType=Mapping)
    n_flow_maps = 0
    for el in root.iter():
        if _xtag(el) != "serviceTask":
            continue
        for child in el:
            if _xtag(child) == "extensionElements":
                props = _read_props(child)
                if props.get("activityType", "") == "Mapping":
                    n_flow_maps += 1
                break

    # 7. Tipo de integración ------------------------------------------
    tipo_integracion = "asíncrona" if async_signals else "síncrona"

    # 8. Adapter types únicos -----------------------------------------
    adapter_types = sorted(set(a["tipo"] for a in adapters_raw))

    # Agrupar por tipo: contar instancias y listar addresses únicos
    adapters_por_tipo: dict[str, dict] = {}
    for a in adapters_raw:
        t = a["tipo"]
        if t not in adapters_por_tipo:
            adapters_por_tipo[t] = {"instancias": 0, "addresses": []}
        adapters_por_tipo[t]["instancias"] += 1
        if a["address"] and a["address"] not in adapters_por_tipo[t]["addresses"]:
            adapters_por_tipo[t]["addresses"].append(a["address"])

    adapters_resumen = [
        {
            "tipo":       t,
            "instancias": adapters_por_tipo[t]["instancias"],
            "addresses":  adapters_por_tipo[t]["addresses"],   # puede ser [] si son params externalizables
        }
        for t in sorted(adapters_por_tipo)
    ]

    # 9. Endpoints externos -------------------------------------------
    endpoints_list = [
        {"adapter": a["tipo"], "url": a["address"]}
        for a in adapters_raw
        if a["tipo"] not in _INTERNAL_ADAPTERS and a["address"]
    ]

    # 10. Complejidad de mappings -------------------------------------
    # Preferir archivos .mmap/.xsl; si no hay, usar count del flujo
    n_maps = len(mapping_files) if mapping_files else n_flow_maps
    if n_maps == 0:
        map_complejidad = "ninguno"
    elif n_maps == 1:
        map_complejidad = f"baja (1 mapping)"
    elif n_maps <= 3:
        map_complejidad = f"media ({n_maps} mappings)"
    else:
        map_complejidad = f"alta ({n_maps} mappings)"

    # 11. Análisis de scripts -----------------------------------------
    scripts_detalle = []
    for sf in script_files:
        lines  = [l for l in sf["contenido"].split("\n") if l.strip()]
        n      = len(lines)
        comp_s = "alta" if n > 100 else ("media" if n > 30 else "baja")
        scripts_detalle.append({
            "nombre":        sf["nombre"],
            "lenguaje":      "Groovy" if sf["nombre"].endswith(".groovy") else "JavaScript",
            "lineas_codigo": n,
            "complejidad":   comp_s,
        })

    # 12. Resultado final ---------------------------------------------
    profile = {
        "iflow_id":         iflow_id,
        "tipo_integracion": tipo_integracion,
        "motivos_async":    list(set(async_signals)),   # vacío = síncrono
        "adapters":         adapter_types,              # lista simple de tipos únicos
        "adapters_detalle": adapters_resumen,           # con instancias y addresses
        "endpoints": {
            "cantidad": len(endpoints_list),
            "lista":    endpoints_list,
        },
        "mappings": {
            "cantidad":    n_maps,
            "complejidad": map_complejidad,
            "archivos":    mapping_files,
        },
        "scripts": {
            "cantidad": len(script_files),
            "detalle":  scripts_detalle,
        },
        "dependencias": {
            "cantidad": len(dependencies),
            "lista":    dependencies,
        },
    }

    return json.dumps(profile, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------
# Tool: detectar anti-patrones en un iFlow
# ------------------------------------------------------------------

# Regex compilados una sola vez (módulo-level para rendimiento)
_AP_RE_URL        = re.compile(r'["\']https?://[^"\'<>\s]{8,}["\']', re.IGNORECASE)
_AP_RE_IP         = re.compile(r'["\'](\d{1,3}\.){3}\d{1,3}(:\d+)?["\']')
_AP_RE_CREDS      = re.compile(
    r'(?i)(password|passwd|pwd|secret|apikey|api_key|token|credential)\s*[=:]\s*["\'][^"\']{3,}["\']'
)
_AP_RE_SLEEP      = re.compile(r'\bThread\.sleep\s*\(', re.IGNORECASE)
_AP_RE_SYSOUT     = re.compile(r'\bSystem\.out\.print', re.IGNORECASE)
_AP_RE_TRY        = re.compile(r'\btry\s*\{')
_AP_RE_COMMENT    = re.compile(r'^\s*(//|/\*|\*)')
_AP_RE_BASIC_AUTH = re.compile(r'Basic\s+[A-Za-z0-9+/=]{10,}')
_AP_RE_BEARER     = re.compile(r'Bearer\s+[A-Za-z0-9._\-]{20,}')

_AP_ADDR_KEYS = ("address", "private.address", "httpAddress",
                 "Private.httpAddress", "wsdlUrl", "serviceUrl",
                 "targetUrl", "wssdUrl", "soapWsdlURL")

_AP_DEPRECATED_ADAPTERS = {
    "FTP":  "FTP transmite credenciales en texto plano. Migrar a SFTP.",
    "XI":   "El adapter XI es legacy. Migrar a HTTP o IDoc según el caso.",
    "LDAP": "Verificar que la conexión use LDAPS (puerto 636) para cifrado TLS.",
}

_AP_SEVERITY_WEIGHT = {"CRITICO": 25, "ALTO": 10, "MEDIO": 5, "BAJO": 2}


def tool_detect_antipatterns(iflow_id: str) -> str:
    """
    Descarga un iFlow de SAP CPI y analiza su XML, scripts y mappings
    en busca de anti-patrones comunes que degradan la calidad, seguridad
    o mantenibilidad del iFlow.

    Detecta 15 categorías de problemas y genera un score de calidad (0-100)
    con recomendaciones concretas para cada hallazgo.
    """
    # 1. Descargar ZIP ------------------------------------------------
    try:
        zip_bytes = get_client().download_iflow(iflow_id)
    except Exception as e:
        return json.dumps({"error": f"No se pudo descargar el iFlow: {e}"})

    # 2. Extraer archivos relevantes ----------------------------------
    iflw_xml      = None
    script_files  = []   # [{nombre, contenido}]
    mapping_files = []   # [{nombre, contenido}]

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        for name in zf.namelist():
            lower = name.lower()
            if lower.endswith(".iflw"):
                iflw_xml = zf.read(name).decode("utf-8", errors="replace")
            elif lower.endswith(".groovy") or (lower.endswith(".js") and not lower.endswith(".json")):
                content = zf.read(name).decode("utf-8", errors="replace")
                script_files.append({"nombre": name.split("/")[-1], "contenido": content})
            elif lower.endswith((".xsl", ".xslt")):
                content = zf.read(name).decode("utf-8", errors="replace")
                mapping_files.append({"nombre": name.split("/")[-1], "contenido": content})
            elif lower.endswith(".mmap"):
                mapping_files.append({"nombre": name.split("/")[-1], "contenido": ""})

    if not iflw_xml:
        return json.dumps({"error": "No se encontró el archivo .iflw en el ZIP"})

    # 3. Parsear XML --------------------------------------------------
    try:
        root = ET.fromstring(iflw_xml.encode("utf-8"))
    except ET.ParseError as e:
        return json.dumps({"error": f"Error al parsear XML: {e}"})

    # 4. Helper para acumular hallazgos --------------------------------
    antipatrones = []

    def add(ap_id, nombre, severidad, descripcion, recomendacion,
            evidencia="", archivo="", linea=None):
        ap = {
            "id":           ap_id,
            "nombre":       nombre,
            "severidad":    severidad,
            "descripcion":  descripcion,
            "recomendacion": recomendacion,
        }
        if evidencia:         ap["evidencia"] = evidencia
        if archivo:           ap["archivo"]   = archivo
        if linea is not None: ap["linea"]     = linea
        antipatrones.append(ap)

    # ── AP001: Sin Exception Subprocess ──────────────────────────────
    # En CPI: subProcess[triggeredByEvent="true"] con errorEventDefinition hijo
    has_exc = False
    for el in root.iter():
        tag = _xtag(el)
        if tag == "subProcess" and el.get("triggeredByEvent", "").lower() == "true":
            for desc in el.iter():
                if _xtag(desc) == "errorEventDefinition":
                    has_exc = True
                    break
        elif tag == "boundaryEvent":
            for child in el:
                if _xtag(child) == "errorEventDefinition":
                    has_exc = True
                    break
        if has_exc:
            break

    if not has_exc:
        add(
            "AP001",
            "Sin Exception Subprocess configurado",
            "ALTO",
            "El iFlow no tiene un Exception Subprocess con Error Start Event. "
            "Los errores en tiempo de ejecución no se capturan y pueden causar "
            "pérdida silenciosa de mensajes.",
            "Agregar un Exception Subprocess (Event-Based) con un Error Start Event. "
            "Dentro, incluir al menos un paso de notificación (email, log) y opcionalmente "
            "reenviar el mensaje fallido a una cola de error.",
        )

    # ── AP002: Endpoints hardcodeados en adapters ─────────────────────
    seen_hardcoded  = set()
    _rag_adapters   = set()   # colectar para la query RAG al final
    for mf in root.iter():
        if _xtag(mf) != "messageFlow":
            continue
        adapter_name = mf.get("name", "").strip()
        if adapter_name:
            _rag_adapters.add(adapter_name)
        props = {}
        for child in mf:
            if _xtag(child) == "extensionElements":
                props = _read_props(child)
                break
        for key in _AP_ADDR_KEYS:
            val = props.get(key, "").strip()
            if not val:
                continue
            # Es un valor hardcodeado si NO empieza con {{ (param externalizable)
            # y NO es un path relativo (no comienza con /)
            is_param    = val.startswith("{{")
            is_relative = val.startswith("/")
            if not is_param and not is_relative and ("." in val or ":" in val):
                dedup_key = f"{adapter_name}:{key}:{val}"
                if dedup_key not in seen_hardcoded:
                    seen_hardcoded.add(dedup_key)
                    add(
                        "AP002",
                        f"Endpoint hardcodeado en adapter '{adapter_name}'",
                        "ALTO",
                        f"El campo '{key}' del adapter '{adapter_name}' contiene un valor "
                        f"literal en lugar de un parámetro externalizable. Esto impide cambiar "
                        f"el endpoint entre ambientes sin redesplegar.",
                        "Reemplazar el valor literal con un parámetro externalizable usando "
                        "la sintaxis {{NombreParametro}}. Configurar el valor en cada ambiente "
                        "mediante Externalized Parameters.",
                        evidencia=f"{key} = \"{val}\"",
                    )

    # ── AP003: iFlow excesivamente complejo ───────────────────────────
    _STEP_TAGS = {"serviceTask", "callActivity", "scriptTask",
                  "userTask", "sendTask", "receiveTask", "subProcess"}
    total_steps = sum(1 for el in root.iter() if _xtag(el) in _STEP_TAGS)
    if total_steps > 20:
        add(
            "AP003",
            f"iFlow excesivamente complejo ({total_steps} pasos)",
            "MEDIO",
            f"El iFlow tiene {total_steps} pasos de procesamiento. Un iFlow con más de 20 "
            "pasos es difícil de entender, mantener y debuggear.",
            "Dividir en sub-iFlows más pequeños y especializados, orquestados mediante "
            "ProcessDirect o JMS. Aplicar el principio de responsabilidad única.",
            evidencia=f"{total_steps} pasos en el flujo",
        )

    # ── AP004: Adapters deprecated o inseguros ────────────────────────
    seen_deprecated = set()
    for mf in root.iter():
        if _xtag(mf) != "messageFlow":
            continue
        adapter_name = mf.get("name", "").strip()
        if adapter_name in _AP_DEPRECATED_ADAPTERS and adapter_name not in seen_deprecated:
            seen_deprecated.add(adapter_name)
            add(
                "AP004",
                f"Adapter deprecated/inseguro: {adapter_name}",
                "MEDIO",
                f"Se detectó el uso del adapter '{adapter_name}', que está deprecated "
                "o presenta problemas de seguridad conocidos.",
                _AP_DEPRECATED_ADAPTERS[adapter_name],
                evidencia=f"messageFlow name=\"{adapter_name}\"",
            )

    # ── AP005: Exceso de Groovy scripts ───────────────────────────────
    n_scripts = len(script_files)
    if n_scripts > 5:
        add(
            "AP005",
            f"Exceso de Groovy/JS scripts ({n_scripts})",
            "ALTO",
            f"El iFlow contiene {n_scripts} scripts. Un número elevado indica que demasiada "
            "lógica de negocio vive en código Groovy/JS, que es difícil de testear, versionar "
            "y mantener por equipos no-developer.",
            "Revisar si alguna lógica puede reemplazarse con steps nativos de CPI: "
            "Content Modifier, Message Mapping, Filter, Splitter. Conservar scripts solo "
            "para lógica que no puede implementarse de otra forma.",
            evidencia=f"{n_scripts} archivos de script",
        )
    elif n_scripts > 3:
        add(
            "AP005",
            f"Varios Groovy/JS scripts ({n_scripts})",
            "MEDIO",
            f"El iFlow tiene {n_scripts} scripts. Evaluar si todos son necesarios.",
            "Revisar si algún script puede consolidarse con otro o reemplazarse "
            "con pasos estándar de CPI.",
            evidencia=f"{n_scripts} archivos de script",
        )

    # ── AP006-AP012: Análisis por script ─────────────────────────────
    for sf in script_files:
        nombre   = sf["nombre"]
        contenido = sf["contenido"]
        lines    = contenido.split("\n")
        code_lines = [l for l in lines if l.strip() and not _AP_RE_COMMENT.match(l)]
        n_lines  = len(code_lines)

        # AP006: God Script
        if n_lines > 200:
            add(
                "AP006",
                f"God Script: {nombre}",
                "ALTO",
                f"'{nombre}' tiene {n_lines} líneas de código efectivo. Concentra demasiada "
                "lógica en un único archivo, dificultando pruebas y mantenimiento.",
                "Dividir en funciones o scripts especializados. Extraer constantes a "
                "parámetros externalizables. Considerar mover lógica a Message Mappings.",
                evidencia=f"{n_lines} líneas de código efectivo",
                archivo=nombre,
            )
        elif n_lines > 100:
            add(
                "AP006",
                f"Script largo: {nombre}",
                "MEDIO",
                f"'{nombre}' tiene {n_lines} líneas de código. Revisar si puede simplificarse.",
                "Evaluar si parte de la lógica puede moverse a pasos estándar de CPI "
                "o extraerse a funciones reutilizables.",
                evidencia=f"{n_lines} líneas de código",
                archivo=nombre,
            )

        # AP007: URL hardcodeada en script
        for i, line in enumerate(lines, 1):
            if _AP_RE_COMMENT.match(line):
                continue
            m = _AP_RE_URL.search(line)
            if m:
                add(
                    "AP007",
                    f"URL hardcodeada en {nombre}",
                    "ALTO",
                    f"Se detectó una URL literal en '{nombre}' (línea {i}). "
                    "Las URLs hardcodeadas hacen el iFlow no portable entre ambientes.",
                    "Leer la URL desde un parámetro externalizable: "
                    "message.getProperty('endpointUrl') o desde el header de la llamada. "
                    "Configurar el valor por ambiente mediante Externalized Parameters.",
                    evidencia=line.strip()[:120],
                    archivo=nombre,
                    linea=i,
                )
                break   # un reporte por script

        # AP008: IP hardcodeada en script
        for i, line in enumerate(lines, 1):
            if _AP_RE_COMMENT.match(line):
                continue
            m = _AP_RE_IP.search(line)
            if m:
                add(
                    "AP008",
                    f"Dirección IP hardcodeada en {nombre}",
                    "CRITICO",
                    f"Se detectó una dirección IP literal en '{nombre}' (línea {i}). "
                    "Las IPs hardcodeadas rompen el iFlow si cambia la infraestructura.",
                    "Nunca hardcodear IPs. Usar nombre de host con DNS, "
                    "parámetros externalizables o Secure Parameter Store de CPI.",
                    evidencia=line.strip()[:120],
                    archivo=nombre,
                    linea=i,
                )
                break

        # AP009: Credenciales hardcodeadas en script
        for i, line in enumerate(lines, 1):
            if _AP_RE_COMMENT.match(line):
                continue
            m = _AP_RE_CREDS.search(line)
            if m:
                add(
                    "AP009",
                    f"Credencial hardcodeada en {nombre}",
                    "CRITICO",
                    f"Posible credencial o secret hardcodeado en '{nombre}' (línea {i}). "
                    "Las credenciales en código fuente representan un riesgo de seguridad crítico.",
                    "URGENTE: Remover la credencial del código inmediatamente. "
                    "Usar Secure Parameter Store de SAP CPI o un alias de credenciales "
                    "(Security Material) y acceder mediante la API de CPI.",
                    evidencia=f"[CENSURADO - credencial detectada en línea {i}]",
                    archivo=nombre,
                    linea=i,
                )
                break

        # AP010: Falta de try/catch en scripts largos
        if not _AP_RE_TRY.search(contenido) and n_lines > 20:
            add(
                "AP010",
                f"Sin manejo de errores en {nombre}",
                "ALTO",
                f"'{nombre}' ({n_lines} líneas) no tiene bloques try/catch. "
                "Una excepción no capturada puede interrumpir el procesamiento "
                "y no ofrece información útil para el diagnóstico.",
                "Envolver la lógica principal en try/catch. En el catch, "
                "asignar un mensaje de error descriptivo al mensaje CPI: "
                "message.setProperty('errorDetail', e.getMessage()) y relanzar "
                "la excepción para que el Exception Subprocess la capture.",
                archivo=nombre,
            )

        # AP011: System.out.println en lugar de logger
        n_sysout = len(_AP_RE_SYSOUT.findall(contenido))
        if n_sysout:
            add(
                "AP011",
                f"System.out.println en {nombre}",
                "BAJO",
                f"Se encontraron {n_sysout} llamadas a System.out.println en '{nombre}'. "
                "En SAP CPI el stdout no es visible en los logs de monitoreo ni de trace.",
                "Reemplazar con el logger de CPI: usar def message = binding.variables['message'] "
                "y loggear propiedades del mensaje. Para debug temporario, "
                "usar message.setHeader('debug', valor) y eliminarlo en producción.",
                evidencia=f"{n_sysout} ocurrencias de System.out.println",
                archivo=nombre,
            )

        # AP012: Thread.sleep
        if _AP_RE_SLEEP.search(contenido):
            add(
                "AP012",
                f"Thread.sleep en {nombre}",
                "ALTO",
                f"Se detectó Thread.sleep en '{nombre}'. Bloquea el thread de procesamiento "
                "de SAP CPI, consumiendo recursos del tenant y pudiendo causar timeouts "
                "en mensajes de larga duración.",
                "Eliminar Thread.sleep. Si se necesita esperar un proceso externo, "
                "usar un adapter con configuración de retry/polling, o rediseñar "
                "el flujo con JMS y un timer que reintente periódicamente.",
                archivo=nombre,
            )

    # ── AP013: Exceso de mappings ─────────────────────────────────────
    n_maps = len(mapping_files)
    if n_maps > 5:
        add(
            "AP013",
            f"Exceso de mappings ({n_maps})",
            "MEDIO",
            f"El iFlow tiene {n_maps} archivos de mapping. Un número elevado puede indicar "
            "transformaciones fragmentadas que podrían consolidarse.",
            "Revisar si mappings consecutivos pueden unificarse en uno solo. "
            "Múltiples transformaciones en cadena impactan performance "
            "y dificultan el entendimiento del flujo de datos.",
            evidencia=f"{n_maps} archivos de mapping",
        )

    # ── AP014: Mapping XSLT trivial (solo identity transform) ─────────
    for mf in mapping_files:
        content = mf.get("contenido", "")
        if not content:
            continue
        # Detectar XSLT con solo copy-of y muy pocas instrucciones reales
        n_xsl_elements = content.count("<xsl:")
        if "xsl:copy-of" in content and n_xsl_elements < 5:
            add(
                "AP014",
                f"Mapping posiblemente trivial: {mf['nombre']}",
                "BAJO",
                f"'{mf['nombre']}' parece ser una transformación de identidad (xsl:copy-of) "
                "sin lógica real de transformación.",
                "Verificar si este mapping es realmente necesario. Una copia simple "
                "puede no requerir un archivo XSLT dedicado y podría eliminarse "
                "o reemplazarse con un Content Modifier.",
                archivo=mf["nombre"],
            )

    # ── AP015: Token/credencial embebida en configuración XML ─────────
    seen_auth = set()
    for mf in root.iter():
        if _xtag(mf) != "messageFlow":
            continue
        adapter_name = mf.get("name", "").strip()
        props = {}
        for child in mf:
            if _xtag(child) == "extensionElements":
                props = _read_props(child)
                break
        for key, val in props.items():
            if (_AP_RE_BASIC_AUTH.search(val) or _AP_RE_BEARER.search(val)):
                dedup_key = f"{adapter_name}:{key}"
                if dedup_key not in seen_auth:
                    seen_auth.add(dedup_key)
                    add(
                        "AP015",
                        f"Credencial embebida en configuración del adapter '{adapter_name}'",
                        "CRITICO",
                        f"El adapter '{adapter_name}' tiene un token o credencial "
                        "embebida directamente en la configuración XML del iFlow. "
                        "Esto expone la credencial en el código fuente del iFlow.",
                        "Usar un alias de credenciales (Security Material > Credentials) "
                        "y referenciar el alias en el campo 'Credential Name' del adapter. "
                        "Nunca embeber tokens o credenciales en la configuración del iFlow.",
                        evidencia="[CENSURADO - credencial detectada en config XML]",
                    )

    # ── Calcular score de calidad ────────────────────────────────────
    penalty = sum(_AP_SEVERITY_WEIGHT.get(a["severidad"], 0) for a in antipatrones)
    score   = max(0, 100 - penalty)

    counts = {"CRITICO": 0, "ALTO": 0, "MEDIO": 0, "BAJO": 0}
    for a in antipatrones:
        counts[a["severidad"]] = counts.get(a["severidad"], 0) + 1

    if score >= 90:
        calificacion = "EXCELENTE"
    elif score >= 75:
        calificacion = "BUENO"
    elif score >= 60:
        calificacion = "MEJORABLE"
    elif score >= 40:
        calificacion = "DEFICIENTE"
    else:
        calificacion = "CRITICO"

    # ── RAG context: iFlows similares del tenant ─────────────────────
    _rag_query    = " ".join([iflow_id] + sorted(_rag_adapters))
    rag_contexto  = _get_rag_context(_rag_query, exclude_id=iflow_id, n=2)

    return json.dumps({
        "iflow_id":  iflow_id,
        "resumen": {
            "total_antipatrones": len(antipatrones),
            "criticos":      counts["CRITICO"],
            "altos":         counts["ALTO"],
            "medios":        counts["MEDIO"],
            "bajos":         counts["BAJO"],
            "score_calidad": score,
            "calificacion":  calificacion,
        },
        "antipatrones": antipatrones,
        "rag_contexto": rag_contexto,
    }, ensure_ascii=False, indent=2)


# ==================================================================
# RAG — helpers y tools
# ==================================================================

def _get_rag_context(query_text: str, exclude_id: str = "", n: int = 2) -> dict:
    """
    Consulta el RAG ChromaDB para obtener iFlows similares al texto dado.
    Retorna silenciosamente un dict con 'nota' si el RAG no está disponible
    o vacío, sin lanzar excepciones (no rompe el flujo de otras tools).
    Usa top 2 por defecto para mantener el contexto compacto.
    """
    try:
        from rag_manager import get_rag
        rag = get_rag()
        if rag.count() == 0:
            return {"nota": "RAG no indexado. Usar la tool 'regenerar_rag' primero."}
        # Pedir 1 extra por si el iFlow actual está en el índice
        similares = rag.get_similar(query_text, n=n + 1)
        # Excluir el propio iFlow del resultado
        if exclude_id:
            similares = [s for s in similares if s["iflow_id"] != exclude_id]
        similares = similares[:n]
        # Devolver solo campos esenciales para mantener el contexto compacto
        compact = [
            {
                "iflow_id":     s["iflow_id"],
                "name":         s["name"],
                "package_name": s["package_name"],
                "adapter_types": s["adapter_types"],
                "n_scripts":    s["n_scripts"],
                "n_mappings":   s["n_mappings"],
                "distancia":    s["distancia"],
            }
            for s in similares
        ]
        return {
            "iflows_similares": compact,
            "total_indexados":  rag.count(),
            "nota": (
                f"Top {len(compact)} iFlows similares en el tenant "
                f"({rag.count()} indexados). Pueden servir de referencia."
            ),
        }
    except Exception as e:
        return {"nota": f"RAG no disponible: {e}"}


# ------------------------------------------------------------------
# Tool: regenerar el índice RAG con todos los iFlows del tenant
# ------------------------------------------------------------------
def tool_regenerate_rag() -> str:
    """
    Descarga TODOS los iFlows del tenant SAP CPI, extrae su contenido
    y reconstruye el índice vectorial ChromaDB desde cero.

    ⚠️  Puede tardar varios minutos dependiendo de la cantidad de iFlows.
    Primera ejecución también descarga el modelo de embedding (~30 MB).
    """
    from rag_manager import get_rag
    rag = get_rag()

    # 1. Limpiar índice anterior
    rag.clear()

    # 2. Obtener todos los iFlows del tenant
    try:
        all_iflows = get_client().filter_iflows("")   # sin filtro = todos
    except Exception as e:
        return json.dumps({"error": f"No se pudo listar los iFlows: {e}"})

    total   = len(all_iflows)
    indexed = 0
    errors  = []

    # 3. Indexar cada iFlow
    for iflow in all_iflows:
        iflow_id   = iflow.get("Id",           "")
        name       = iflow.get("Name",         "")
        pkg_id     = iflow.get("_PackageId",   "")
        pkg_name   = iflow.get("_PackageName", "")
        description = iflow.get("Description", "") or ""

        try:
            zip_bytes = get_client().download_iflow(iflow_id)
        except Exception as e:
            errors.append({"iflow_id": iflow_id, "error": str(e)})
            continue

        # Extraer contenido del ZIP
        adapter_types = []
        step_names    = []
        script_names  = []
        script_texts  = []
        mapping_names = []

        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                for fname in zf.namelist():
                    lower = fname.lower()
                    if lower.endswith(".iflw"):
                        xml_text = zf.read(fname).decode("utf-8", errors="replace")
                        try:
                            xml_root = ET.fromstring(xml_text.encode("utf-8"))
                            # Adapters desde messageFlow[name]
                            for mf_el in xml_root.iter():
                                if _xtag(mf_el) == "messageFlow":
                                    aname = mf_el.get("name", "").strip()
                                    if aname and aname not in adapter_types:
                                        adapter_types.append(aname)
                            # Nombres de pasos relevantes
                            for el in xml_root.iter():
                                if _xtag(el) in ("serviceTask", "callActivity"):
                                    sname = el.get("name", "").strip()
                                    if sname and len(sname) > 2:
                                        step_names.append(sname)
                        except Exception:
                            pass
                    elif lower.endswith(".groovy") or (
                            lower.endswith(".js") and not lower.endswith(".json")):
                        bname = fname.split("/")[-1]
                        script_names.append(bname)
                        content = zf.read(fname).decode("utf-8", errors="replace")
                        script_texts.append(content[:400])
                    elif lower.endswith((".mmap", ".xsl", ".xslt")):
                        mapping_names.append(fname.split("/")[-1])
        except Exception as e:
            errors.append({"iflow_id": iflow_id, "error": f"ZIP error: {e}"})
            continue

        # Construir texto del documento para embedding
        doc_parts = [
            f"{name} — {pkg_name}",
            f"Descripción: {description[:200]}" if description else "",
            f"Adapters: {', '.join(adapter_types)}"  if adapter_types  else "",
            f"Pasos: {', '.join(step_names[:10])}"   if step_names     else "",
            f"Scripts: {', '.join(script_names)}"    if script_names   else "",
            f"Mappings: {', '.join(mapping_names)}"  if mapping_names  else "",
        ]
        if script_texts:
            combined = " | ".join(script_texts)
            doc_parts.append(f"Contenido scripts: {combined[:800]}")

        document = "\n".join(p for p in doc_parts if p)

        metadata = {
            "iflow_id":    iflow_id,
            "name":        name,
            "package_id":  pkg_id,
            "package_name": pkg_name,
            "description": description[:200],
            "adapter_types": ", ".join(adapter_types),
            "n_scripts":   len(script_names),
            "n_mappings":  len(mapping_names),
            "indexed_at":  datetime.now().isoformat(),
        }

        try:
            rag.upsert(iflow_id, document, metadata)
            indexed += 1
        except Exception as e:
            errors.append({"iflow_id": iflow_id, "error": f"Upsert error: {e}"})

    return json.dumps({
        "status":       "ok",
        "total_iflows": total,
        "indexed":      indexed,
        "errors":       len(errors),
        "error_details": errors[:10],
        "timestamp":    datetime.now().isoformat(),
        "nota": (
            f"RAG reconstruido: {indexed}/{total} iFlows indexados. "
            "Ahora las tools de análisis incluirán contexto del tenant."
        ),
    }, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------
# Tool: buscar en el RAG
# ------------------------------------------------------------------
def tool_query_rag(query: str, n_results: int = 5) -> str:
    """
    Realiza una búsqueda semántica en el índice RAG y retorna los iFlows
    más similares a la consulta dada. Útil para encontrar iFlows de referencia
    antes de analizar o para explorar el tenant.
    """
    from rag_manager import get_rag
    rag = get_rag()

    if rag.count() == 0:
        return json.dumps({
            "error": "El índice RAG está vacío. Ejecutar 'regenerar_rag' primero."
        })

    n        = max(1, min(n_results, 20))
    similares = rag.get_similar(query, n=n)

    return json.dumps({
        "query":            query,
        "total_indexados":  rag.count(),
        "n_resultados":     len(similares),
        "resultados":       similares,
    }, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------
# Helper: llamada a Claude con retry automático en 429 (rate limit)
# ------------------------------------------------------------------
def _call_claude_with_retry(client, model: str, max_tokens: int, messages: list,
                             max_retries: int = 3, wait_seconds: int = 60):
    """
    Llama a client.messages.create() con retry automático ante error 429.
    Espera `wait_seconds` segundos entre intentos. Máximo `max_retries` intentos.
    Lanza la excepción original si se agotan los reintentos.
    """
    import time as _time
    import anthropic as _anthropic

    for attempt in range(1, max_retries + 1):
        try:
            return client.messages.create(
                model=model,
                max_tokens=max_tokens,
                messages=messages,
            )
        except _anthropic.RateLimitError as e:
            if attempt < max_retries:
                print(f"[generate_iflow] 429 Rate Limit en intento {attempt}/{max_retries}. "
                      f"Esperando {wait_seconds}s antes de reintentar...")
                _time.sleep(wait_seconds)
            else:
                print(f"[generate_iflow] 429 Rate Limit — se agotaron los {max_retries} intentos.")
                raise


# ------------------------------------------------------------------
# Helper interno: armar el ZIP con la estructura correcta de SAP CPI
# ------------------------------------------------------------------
def _build_iflow_zip(
    iflow_xml:    str,
    iflow_name:   str,
    iflow_id:     str = "",
    scripts:      dict = None,
    description:  str = "",
    package_id:   str = "",
    package_name: str = "",
) -> tuple:
    """
    Empaqueta un iFlow en un ZIP con la estructura exacta que espera SAP CPI para import.

    Estructura generada:
      META-INF/MANIFEST.MF
      metainfo.prop
      src/main/resources/scenarioflows/integrationflow/<id>.iflw
      src/main/resources/parameters.prop
      src/main/resources/parameters.propdef
      src/main/resources/script/<nombre>.groovy   (si hay scripts)
      .project

    Retorna: (zip_path, zip_filename, metadata_dict)
    """
    scripts = scripts or {}

    # Normalizar IDs y nombres
    if not iflow_id:
        iflow_id = re.sub(r"[^\w\-.]", "_", iflow_name)[:60]
    safe_id   = re.sub(r"[^\w\-.]", "_", iflow_id)[:60]
    safe_name = iflow_name.strip() or safe_id

    # Extraer parámetros externalizables {{ParamName}} del XML
    param_names = sorted(set(re.findall(r"\{\{(\w+)\}\}", iflow_xml)))

    ts       = datetime.now()
    ts_prop  = ts.strftime("%a %b %d %H:%M:%S UTC %Y")   # "Tue Mar 24 14:00:00 UTC 2026"
    ts_file  = ts.strftime("%Y%m%d_%H%M%S")

    # ── META-INF/MANIFEST.MF ──────────────────────────────────────
    manifest = (
        f"Bundle-Name: {safe_name}\n"
        f"Bundle-SymbolicName: {safe_id}\n"
        f"Bundle-Version: 1.0.0\n"
        f"Bundle-ManifestVersion: 2\n"
    )

    # ── metainfo.prop ─────────────────────────────────────────────
    metainfo = (
        f"#{ts_prop}\n"
        f"category=\n"
        f"cpi.sdk.version=\n"
        f"display_name={safe_name}\n"
        f"id={safe_id}\n"
        f"packageId={package_id}\n"
        f"packageTechnicalName={package_name}\n"
        f"version=1.0.0\n"
        f"vendor=\n"
        f"description={description}\n"
    )

    # ── parameters.prop ───────────────────────────────────────────
    param_prop_lines = [f"#{ts_prop}"]
    for p in param_names:
        param_prop_lines.append(f"{p}=")
    parameters_prop = "\n".join(param_prop_lines) + "\n"

    # ── parameters.propdef ────────────────────────────────────────
    param_entries = ""
    for p in param_names:
        param_entries += (
            f"    <map:parameter>\n"
            f"        <map:name>{p}</map:name>\n"
            f"        <map:value/>\n"
            f"        <map:type>xs:string</map:type>\n"
            f"        <map:sensitive>false</map:sensitive>\n"
            f"    </map:parameter>\n"
        )
    parameters_propdef = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<map:parameter-to-mapping-table '
        'xmlns:map="http://sap.hana.com/pi/pcp/xi/parameter/1.0">\n'
        f'{param_entries}'
        '</map:parameter-to-mapping-table>\n'
    )

    # ── .project ──────────────────────────────────────────────────
    dot_project = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<projectDescription>\n'
        f'\t<name>{safe_id}</name>\n'
        '\t<comment></comment>\n'
        '\t<projects>\n'
        '\t</projects>\n'
        '\t<buildSpec>\n'
        '\t</buildSpec>\n'
        '\t<natures>\n'
        '\t</natures>\n'
        '</projectDescription>\n'
    )

    # ── Empaquetar ────────────────────────────────────────────────
    zip_filename  = f"{safe_id}_{ts_file}.zip"
    generated_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_iflows")
    os.makedirs(generated_dir, exist_ok=True)
    zip_path = os.path.join(generated_dir, zip_filename)

    iflw_path_in_zip = (
        f"src/main/resources/scenarioflows/integrationflow/{safe_id}.iflw"
    )

    # Normalizar nombres de scripts una sola vez (para ZIP y para estructura)
    safe_scripts = {re.sub(r"[^\w.]", "_", sname): scontent
                    for sname, scontent in scripts.items()}

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("META-INF/MANIFEST.MF",           manifest)
        zf.writestr("metainfo.prop",                   metainfo)
        zf.writestr(iflw_path_in_zip,                  iflow_xml)
        zf.writestr("src/main/resources/parameters.prop",    parameters_prop)
        zf.writestr("src/main/resources/parameters.propdef", parameters_propdef)
        for safe_sname, scontent in safe_scripts.items():
            zf.writestr(f"src/main/resources/script/{safe_sname}", scontent)
        zf.writestr(".project", dot_project)

    structure = (
        ["META-INF/MANIFEST.MF", "metainfo.prop", iflw_path_in_zip,
         "src/main/resources/parameters.prop",
         "src/main/resources/parameters.propdef"]
        + [f"src/main/resources/script/{s}" for s in safe_scripts]
        + [".project"]
    )

    metadata = {
        "zip_filename":              zip_filename,
        "zip_path":                  zip_path,
        "download_url":              f"/download/generated/{zip_filename}",
        "size_bytes":                os.path.getsize(zip_path),
        "parametros_externalizados": param_names,
        "scripts_incluidos":         list(scripts.keys()),
        "estructura":                structure,
    }
    return zip_path, zip_filename, metadata


# ------------------------------------------------------------------
# Tool: empaquetar un XML en ZIP con la estructura correcta de SAP CPI
# ------------------------------------------------------------------
def tool_generate_iflow_zip(
    iflow_xml:    str,
    iflow_name:   str,
    iflow_id:     str  = "",
    scripts_json: str  = "{}",
    description:  str  = "",
    package_id:   str  = "",
    package_name: str  = "",
) -> str:
    """
    Toma el XML de un iFlow y lo empaqueta en un ZIP con la estructura
    exacta que espera SAP CPI para importar:

      META-INF/MANIFEST.MF
      metainfo.prop
      src/main/resources/scenarioflows/integrationflow/<id>.iflw
      src/main/resources/parameters.prop
      src/main/resources/parameters.propdef
      src/main/resources/script/<nombre>.groovy  (si hay scripts)
      .project

    Uso típico: el usuario genera el XML con generate_iflow y luego llama
    a esta tool para obtener el ZIP importable.
    Retorna el path y la URL de descarga del ZIP generado.
    """
    # Parsear scripts desde JSON string
    try:
        scripts = json.loads(scripts_json) if scripts_json.strip() not in ("", "{}") else {}
        if not isinstance(scripts, dict):
            scripts = {}
    except Exception:
        scripts = {}

    try:
        zip_path, zip_filename, meta = _build_iflow_zip(
            iflow_xml    = iflow_xml,
            iflow_name   = iflow_name,
            iflow_id     = iflow_id,
            scripts      = scripts,
            description  = description,
            package_id   = package_id,
            package_name = package_name,
        )
        return json.dumps({
            "status":       "ok",
            "iflow_name":   iflow_name,
            "iflow_id":     iflow_id or re.sub(r"[^\w\-.]", "_", iflow_name)[:60],
            **meta,
            "instrucciones": (
                "Importar en SAP CPI: "
                "Integration Suite → Design → Import → Integration Flow → seleccionar ZIP."
            ),
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)}, ensure_ascii=False)


# ------------------------------------------------------------------
# Tool: generar un iFlow nuevo desde descripción (usando RAG como contexto)
# ------------------------------------------------------------------
def tool_generate_iflow(description: str, iflow_name: str = "") -> str:
    """
    Busca en el RAG iFlows similares del tenant para usarlos como referencia
    al diseñar un nuevo iFlow. Devuelve contexto estructurado (adapters, scripts,
    patrones) que el agente usa para elaborar su guía paso a paso.
    No genera archivos ni llama a Claude API.
    """
    from rag_manager import get_rag
    rag = get_rag()

    similares = []
    if rag.count() > 0:
        similares = rag.get_similar(description, n=2)

    ctx = []
    for sim in similares:
        ctx.append({
            "iflow_id":      sim.get("iflow_id", ""),
            "name":          sim.get("name", ""),
            "adapter_types": sim.get("adapter_types", []),
            "n_scripts":     sim.get("n_scripts", 0),
            "n_mappings":    sim.get("n_mappings", 0),
            "resumen":       sim.get("resumen", ""),
        })

    return json.dumps({
        "status":         "ok",
        "iflow_buscado":  description,
        "iflows_similares_en_tenant": ctx,
        "nota": (
            "Usá estos iFlows como referencia para la guía. "
            "Respondé con: 1) Enfoque general, 2) Paso a paso en SAP CPI designer."
        ),
    }, ensure_ascii=False, indent=2)
